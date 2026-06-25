"""
Document processor — extracts text from PDF, DOCX, TXT and splits into chunks.
Uses LlamaIndex's node pipeline for smart chunking with metadata preservation.
"""
from typing import List, Tuple, Optional
from pathlib import Path
import hashlib

from llama_index.core import Document as LIDocument
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import TextNode

from app.core.config import settings


def extract_text_pdf(file_path: str) -> Tuple[str, dict]:
    """Extract text from PDF, return (text, metadata)."""
    import pypdf

    reader = pypdf.PdfReader(file_path)
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n".join(pages_text)
    metadata = {
        "page_count": len(reader.pages),
        "title": reader.metadata.get("/Title", ""),
        "author": reader.metadata.get("/Author", ""),
    }
    return full_text, metadata


def extract_text_docx(file_path: str) -> Tuple[str, dict]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    metadata = {
        "paragraph_count": len(paragraphs),
        "title": doc.core_properties.title or "",
        "author": doc.core_properties.author or "",
    }
    return full_text, metadata


def extract_text_txt(file_path: str) -> Tuple[str, dict]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return text, {"char_count": len(text)}


EXTRACTORS = {
    "pdf": extract_text_pdf,
    "docx": extract_text_docx,
    "txt": extract_text_txt,
    "md": extract_text_txt,
}


def extract_text(file_path: str, file_type: str) -> Tuple[str, dict]:
    extractor = EXTRACTORS.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    return extractor(file_path, )


def chunk_document(
    text: str,
    doc_id: str,
    filename: str,
    doc_metadata: dict,
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[TextNode]:
    """
    Split text into overlapping chunks using LlamaIndex SentenceSplitter.
    SentenceSplitter is preferred over RecursiveCharacterTextSplitter because it
    respects sentence boundaries, producing more semantically coherent chunks.
    """
    chunk_size = chunk_size or settings.CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP

    splitter = SentenceSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    li_doc = LIDocument(
        text=text,
        doc_id=doc_id,
        metadata={
            "doc_id": doc_id,
            "filename": filename,
            **doc_metadata,
        },
    )

    nodes = splitter.get_nodes_from_documents([li_doc])

    # Add chunk index for citation tracking
    for i, node in enumerate(nodes):
        node.metadata["chunk_index"] = i
        node.metadata["chunk_hash"] = hashlib.md5(node.text.encode()).hexdigest()[:8]

    return nodes
